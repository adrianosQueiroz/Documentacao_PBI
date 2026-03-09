import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def gerar_documentacao_completa(caminho_excel_input):
    # ==========================================================
    # 📝 CAMPOS PARA PREENCHIMENTO MANUAL (AJUSTE AQUI)
    # ==========================================================
    NOME_PROJETO      = ""
    RESPONSAVEL       = ""
    VERSAO            = "1.0"
    DEPARTAMENTO      = ""
    DESCRICAO_NEGOCIO = (
        "Este dashboard monitora os KPIs de Lead Time desde a entrada "
        "até o faturamento, permitindo a gestão de gargalos operacionais."
    )
    # ==========================================================

    # Definição de nomes de saída
    nome_base = os.path.splitext(os.path.basename(caminho_excel_input))[0]
    diretorio = os.path.dirname(caminho_excel_input)
    saida_excel = os.path.join(diretorio, f"{nome_base}_Documentacao.xlsx")
    saida_word = os.path.join(diretorio, f"{nome_base}_Documentacao.docx")

    # 1. Carregar dados
    df_original = pd.read_excel(caminho_excel_input)
    cols = df_original.columns.tolist()

    # --- PROCESSAMENTO ---
    # Medidas (DAX e Pastas)
    df_medidas = pd.DataFrame()
    if 'Type' in cols:
        df_medidas = df_original[df_original['Type'] == 'Measure'].copy()
        map_med = {'Name': 'Nome', 'Table': 'Tabela', 'Display_folder': 'Pasta', 'Source': 'Cálculo DAX'}
        df_medidas = df_medidas[[c for c in map_med.keys() if c in cols]].drop_duplicates()
        df_medidas.rename(columns=map_med, inplace=True)

    # Tabelas
    df_tabelas = df_original[['Table']].dropna().drop_duplicates() if 'Table' in cols else pd.DataFrame()

    # Colunas (Dicionário de Dados)
    df_colunas = pd.DataFrame()
    if 'Type' in cols:
        df_colunas = df_original[df_original['Type'].isin(['Column', 'Calculated Column'])].copy()
        map_col = {'Table': 'Tabela', 'Name': 'Coluna', 'Data_type': 'Formato'}
        df_colunas = df_colunas[[c for c in map_col.keys() if c in cols]].drop_duplicates()
        df_colunas.rename(columns=map_col, inplace=True)

    # --- SALVAR EXCEL ---
    with pd.ExcelWriter(saida_excel) as writer:
        df_medidas.to_excel(writer, sheet_name='Dicionário de Medidas', index=False)
        df_tabelas.to_excel(writer, sheet_name='Tabelas', index=False)
        df_colunas.to_excel(writer, sheet_name='Dicionário de Colunas', index=False)

    # --- GERAR WORD ---
    doc = Document()
    
    # Capa
    titulo = doc.add_heading('\n\n\n' + NOME_PROJETO, 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_capa = doc.add_paragraph()
    p_capa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_capa.add_run(f"Responsável: {RESPONSAVEL}\nVersão: {VERSAO}\nDepto: {DEPARTAMENTO}")
    run.font.size = Pt(14)
    doc.add_page_break()

    # Seção 1: Objetivo
    doc.add_heading('1. Objetivo do Dashboard', level=1)
    doc.add_paragraph(DESCRICAO_NEGOCIO)

    # Seção 2: Modelo de Dados (ESPAÇO PARA O PRINT)
    doc.add_heading('2. Modelo de Dados (Star Schema)', level=1)
    doc.add_paragraph("A arquitetura de dados segue o modelo abaixo:")
    p_placeholder = doc.add_paragraph()
    p_placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_p = p_placeholder.add_run("\n\n[ 📸 COLE AQUI O PRINT DO SEU MODELO DE DADOS ]\n\n")
    run_p.bold = True
    run_p.font.size = Pt(12)

    # Função auxiliar para tabelas no Word
    def add_table_doc(titulo, df):
        if df.empty: return
        doc.add_heading(titulo, level=1)
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        for i, col in enumerate(df.columns):
            table.rows[0].cells[i].text = col
        for _, row in df.iterrows():
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val) if pd.notna(val) else ""

    # Adicionar as seções técnicas
    add_table_doc("3. Tabelas do Modelo", df_tabelas)
    add_table_doc("4. Dicionário de Medidas", df_medidas)
    add_table_doc("5. Dicionário de Colunas", df_colunas)

    doc.save(saida_word)
    print(f"✅ Documentação completa gerada:\nExcel: {saida_excel}\nWord: {saida_word}")

# Execução
gerar_documentacao_completa("results.xlsx")